from __future__ import annotations

from pathlib import Path

import pandas as pd
import pytesseract
from docx import Document as DocxDocument
from PIL import Image
from pypdf import PdfReader

try:
    import pypdfium2 as pdfium
except ImportError:  # pragma: no cover - exercised through fallback behavior
    pdfium = None


class DocumentParseError(RuntimeError):
    pass


def parse_document(path: str) -> str:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(file_path)
    if suffix in {".docx"}:
        return parse_docx(file_path)
    if suffix in {".txt", ".md"}:
        return file_path.read_text(encoding="utf-8", errors="ignore")
    if suffix in {".png", ".jpg", ".jpeg", ".tiff"}:
        return pytesseract.image_to_string(Image.open(file_path))
    if suffix in {".csv", ".xlsx", ".xls"}:
        return parse_spreadsheet(file_path)
    return file_path.read_text(encoding="utf-8", errors="ignore")


def parse_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    chunks: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        chunks.append(f"[Page {index}]\n{page.extract_text() or ''}")
    extracted_text = "\n\n".join(chunks).strip()
    if has_meaningful_text(extracted_text):
        return extracted_text

    ocr_text = parse_pdf_with_ocr(path)
    if has_meaningful_text(ocr_text):
        return ocr_text

    raise DocumentParseError("PDF text extraction produced no usable text, including OCR fallback.")


def parse_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def parse_spreadsheet(path: Path) -> str:
    if path.suffix.lower() == ".csv":
        data = pd.read_csv(path)
    else:
        data = pd.read_excel(path)
    return data.to_csv(index=False)


def has_meaningful_text(text: str, threshold: int = 24) -> bool:
    return len("".join(text.split())) >= threshold


def parse_pdf_with_ocr(path: Path) -> str:
    if pdfium is None:
        return ""

    chunks: list[str] = []
    pdf = pdfium.PdfDocument(str(path))
    try:
        for index, page in enumerate(pdf, start=1):
            bitmap = page.render(scale=2).to_pil()
            try:
                page_text = pytesseract.image_to_string(bitmap).strip()
            finally:
                bitmap.close()
            if page_text:
                chunks.append(f"[Page {index}]\n{page_text}")
    finally:
        pdf.close()

    return "\n\n".join(chunks)
